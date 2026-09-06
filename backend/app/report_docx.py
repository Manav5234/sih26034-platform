"""DOCX report renderer using python-docx."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.report import ReportData, _format_date_value, _format_value

_DASH = "\u2014"

VERDICT_COLORS = {
    "SATISFIED": RGBColor(0x16, 0xa3, 0x4a),
    "VIOLATION": RGBColor(0xdc, 0x26, 0x26),
    "NOT_VERIFIED": RGBColor(0xd9, 0x77, 0x06),
    "CONFLICT": RGBColor(0x7c, 0x3a, 0xed),
}


def render_docx(report: ReportData) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Legal Metrology Compliance Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Scan ID: {report.scan_id}")
    doc.add_paragraph(f"Generated: {report.generated_at}")
    doc.add_paragraph("")

    # Overall status
    status = report.overall_status or "UNKNOWN"
    p = doc.add_paragraph()
    p.add_run("Overall Status: ").bold = True
    p.add_run(status)

    # Product info
    if report.product:
        doc.add_heading("Product Information", level=1)
        table = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for key, label in [("name", "Name"), ("brand", "Brand"), ("category", "Category"),
                           ("manufacturer", "Manufacturer"), ("barcode", "Barcode"),
                           ("mrp", "MRP"), ("quantity", "Net Quantity")]:
            val = report.product.get(key) or "\u2014"
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(val)
        doc.add_paragraph("")

    # Declarations
    doc.add_heading("Extracted Declarations", level=1)
    for f in report.fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{f.field_name}")
        run.bold = True

        # Show verdict: if officer-corrected, show original AI verdict + officer override
        if f.officer_correction and f.ai_verdict:
            ai_run = p.add_run(f" \u2014 AI: {f.ai_verdict}")
            ai_run.bold = True
            ai_color = VERDICT_COLORS.get(f.ai_verdict)
            if ai_color:
                ai_color = ai_color
                ai_run.font.color.rgb = ai_color
            arrow_run = p.add_run(" \u2192 ")
            officer_run = p.add_run(f"Officer: {f.verdict}")
            officer_run.bold = True
            o_color = VERDICT_COLORS.get(f.verdict)
            if o_color:
                officer_run.font.color.rgb = o_color
        else:
            verdict_run = p.add_run(f" \u2014 {f.verdict}")
            verdict_run.bold = True
            color = VERDICT_COLORS.get(f.verdict)
            if color:
                verdict_run.font.color.rgb = color

        # Nutrition facts: render as table
        if f.field_name == "nutrition_facts" and isinstance(f.extracted_value, list):
            table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
            table.rows[0].cells[0].text = "Nutrient"
            table.rows[0].cells[1].text = "Value"
            table.rows[0].cells[2].text = "Unit"
            table.rows[0].cells[3].text = "Confidence"
            for n in f.extracted_value:
                val = n.get("value")
                row = table.add_row()
                row.cells[0].text = n.get("nutrient", "").replace("_", " ").title()
                row.cells[1].text = f"{val}" if val is not None else "\u2014"
                row.cells[2].text = n.get("unit", "")
                row.cells[3].text = f"{n.get('confidence', 0):.0%}"
        # Date fields: render as readable string
        elif f.field_name in ("manufacture_date", "expiry_date") and isinstance(f.extracted_value, dict):
            doc.add_paragraph(f"Value: {_format_date_value(f.extracted_value)}", style="List Bullet")
        # Cautions: render as quoted text or "Not present"
        elif f.field_name == "cautions":
            if f.extracted_value and isinstance(f.extracted_value, dict) and f.extracted_value.get("present"):
                doc.add_paragraph(f"Value: \"{f.extracted_value.get('text', '')}\"", style="List Bullet")
            else:
                doc.add_paragraph("Value: Not present", style="List Bullet")
        # Default: standard value display
        else:
            doc.add_paragraph(f"Value: {f.display_value}", style="List Bullet")

        # Reason
        if f.officer_correction and f.ai_reason:
            doc.add_paragraph(f"AI Reason: {f.ai_reason}", style="List Bullet")
            doc.add_paragraph(f"Officer Note: {f.reason}", style="List Bullet")
        elif f.reason:
            doc.add_paragraph(f"Reason: {f.reason}", style="List Bullet")

        # Confidence: show "Officer-reviewed" for officer-touched fields
        if f.officer_correction:
            doc.add_paragraph(f"Confidence: Officer-reviewed | Rule: {f.rule_id or _DASH}", style="List Bullet")
        else:
            doc.add_paragraph(f"Confidence: {f.confidence:.1%} | Rule: {f.rule_id or _DASH}", style="List Bullet")

        # Evidence
        for ev in f.evidence:
            doc.add_paragraph(
                f"[{ev['source_label']}] \"{ev.get('raw_text', '')}\" (conf: {ev['confidence']:.0%})",
                style="List Bullet 2"
            )

        # Officer correction
        if f.officer_correction:
            oc = f.officer_correction
            corrected_display = oc.get("corrected_value", "\u2014")
            if isinstance(corrected_display, dict):
                corrected_display = _format_value(f.field_name, corrected_display)
            p = doc.add_paragraph(style="List Bullet 2")
            run = p.add_run(f"Officer Correction: {oc.get('officer_name', 'Officer')} set to {corrected_display} \u2014 {oc.get('reason', '')}")
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

    # Inspections
    if report.inspections:
        doc.add_heading("Officer Reviews", level=1)
        for insp in report.inspections:
            doc.add_paragraph(f"Officer: {insp.officer_name} | {insp.created_at}")
            if insp.location:
                loc = insp.location
                loc_text = f"Location: ({loc.latitude:.6f}, {loc.longitude:.6f})"
                if loc.accuracy_meters:
                    loc_text += f" \u00b1{loc.accuracy_meters:.0f}m"
                loc_text += f" [source: {loc.source}]"
                if loc.address_text:
                    loc_text += f" \u2014 {loc.address_text}"
                doc.add_paragraph(loc_text, style="List Bullet")
            for a in insp.actions:
                doc.add_paragraph(
                    f"{a.get('action', '')}: {a.get('field_name', '')} \u2014 {a.get('reason', '')}",
                    style="List Bullet"
                )
            doc.add_paragraph("")

    # Summary
    doc.add_heading("Summary", level=1)
    table = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
    table.rows[0].cells[0].text = "Verdict"
    table.rows[0].cells[1].text = "Count"
    for label, key in [("Satisfied", "SATISFIED"), ("Violation", "VIOLATION"),
                       ("Not Verified", "NOT_VERIFIED"), ("Conflict", "CONFLICT")]:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(report.summary.get(key, 0))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
